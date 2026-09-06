# -*- coding: utf-8 -*-
"""
LLM(Ollama) 기반 민감정보 마스킹 스크립트
- mask_sensitive.py 에서 정규식으로 처리한 항목 + 문맥 판단이 필요한 항목까지
  로컬 Ollama 모델(qwen3-vl:4b-instruct)에게 판단시켜 '*'로 치환한다.
- Ollama가 로컬에서 실행 중이어야 한다. (기본: http://localhost:11434)
"""
import json
import time

import docx
import requests

INPUT_FILE = "한백시멘트_AI연계_검토보고서.docx"
OUTPUT_FILE = "한백시멘트_AI연계_검토보고서_마스킹_LLM.docx"

OLLAMA_URL = "http://localhost:11434/api/chat"
MODEL_NAME = "qwen3-vl:4b-instruct"

SYSTEM_PROMPT = """\
너는 사내 문서의 민감정보를 마스킹하는 필터다.
아래 규칙에 해당하는 부분만 정확히 찾아 원래 글자 수만큼 '*'로 치환하고,
그 외의 문장/표 내용은 한 글자도 바꾸지 말고 그대로 유지해야 한다.

[마스킹 대상]
1. 회사명: "한백시멘트", "(주)라우트나인" / "㈜라우트나인"
2. 연락처(전화번호), 이메일 주소
3. DB 접속 문자열(예: postgresql://... 형태), IP 주소, API 키 같은 시스템 접속 정보
4. 내부/외부 도메인·호스트명 (예: hanbaek-cement, route9-ai 가 포함된 도메인)
5. 사람 이름(성명)
6. 영업비밀성 수치: 계약금액, 위약금 비율, 인력 감축률, 아직 공시되지 않은
   내부 목표치, 특정 고객사명과 그 고객사와의 공급 물량/단가, 제조원가 관련 수치

[마스킹하지 않는 것]
- 이미 공시자료/홈페이지 등에 공개된 것으로 명시된 수치나 정보
  (예: 생산능력 수치, "공시자료 기준으로 공개 가능"이라고 문서에 적힌 수치)
- 위 규칙에 해당하지 않는 일반 서술, 부서명, 직급명, 표 헤더 등

[출력 형식]
- 반드시 마스킹 처리된 최종 텍스트만 출력한다.
- 설명, 접두사, 따옴표, 코드블록(```) 등 어떤 부가 텍스트도 붙이지 않는다.
- 입력이 비어 있거나 마스킹할 대상이 없으면 입력을 그대로 출력한다.
"""

_cache: dict[str, str] = {}


def llm_mask_text(text: str, retries: int = 3, timeout: int = 60) -> str:
    """Ollama 모델에 텍스트를 보내 마스킹된 결과를 받아온다."""
    if not text or not text.strip():
        return text

    if text in _cache:
        return _cache[text]

    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": text},
        ],
        "stream": False,
        "options": {"temperature": 0},
    }

    last_error = None
    for attempt in range(1, retries + 1):
        try:
            resp = requests.post(OLLAMA_URL, json=payload, timeout=timeout)
            resp.raise_for_status()
            data = resp.json()
            masked = data["message"]["content"].strip()
            # 모델이 실수로 코드블록/따옴표를 씌운 경우 제거
            if masked.startswith("```"):
                masked = masked.strip("`").strip()
            if masked.startswith('"') and masked.endswith('"') and len(masked) >= 2:
                masked = masked[1:-1]
            _cache[text] = masked
            return masked
        except Exception as e:  # noqa: BLE001
            last_error = e
            time.sleep(1)

    print(f"[경고] LLM 호출 실패, 원본 유지: {text[:30]!r} ({last_error})")
    _cache[text] = text
    return text


def mask_paragraph(paragraph):
    for run in paragraph.runs:
        if run.text:
            run.text = llm_mask_text(run.text)


def mask_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                mask_paragraph(paragraph)
            for nested_table in cell.tables:
                mask_table(nested_table)


def main():
    document = docx.Document(INPUT_FILE)

    for paragraph in document.paragraphs:
        mask_paragraph(paragraph)

    for table in document.tables:
        mask_table(table)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            mask_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            mask_paragraph(paragraph)

    document.save(OUTPUT_FILE)
    print(f"완료: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
