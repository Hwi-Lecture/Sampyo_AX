# -*- coding: utf-8 -*-
"""
정규식 기반 민감정보 마스킹 스크립트
1차 대상(범용 규칙): 한백시멘트, (주)라우트나인, 연락처(전화번호), 이메일
2차 대상(문서 특화 규칙 - 추가 검토 항목): DB 접속 문자열/IP, API 키,
        내부 도메인/호스트명, 인명, 영업비밀성 수치
"""
import re
import docx

INPUT_FILE = "한백시멘트_AI연계_검토보고서.docx"
OUTPUT_FILE = "한백시멘트_AI연계_검토보고서_마스킹.docx"

# --- 1차: 범용으로 재사용 가능한 규칙 ---
PATTERNS = [
    re.compile(r"한백시멘트"),                     # 회사명: 한백시멘트
    re.compile(r"(?:\(주\)|㈜)\s*라우트나인"),        # 회사명: (주)라우트나인 / ㈜라우트나인
    re.compile(r"01[016789]-\d{3,4}-\d{4}"),        # 연락처(휴대전화)
    re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"),  # 이메일
]

# --- 2차: 이 문서에 한정된 추가 검토 항목(범용 규칙으로 보기는 어려움) ---
EXTRA_PATTERNS = [
    re.compile(r"postgresql://\S+"),                       # DB 접속 문자열(계정/비밀번호 포함)
    re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b"),             # IP 주소
    re.compile(r"\brt9_live_[0-9a-f]+\b"),                  # 벤더 API 키
    re.compile(r"[\w.-]*(?:hanbaek-cement|route9-ai)[\w.-]*"),  # 내부/외부 도메인·호스트명
    re.compile(r"김도현|박현수|정하윤"),                        # 인명(문서에 등장하는 성명 목록)
    re.compile(r"3억\s*2천만원"),                             # 계약금액
    re.compile(r"(?<!\d)30%(?!\d)"),                         # 위약금 비율
    re.compile(r"(?<!\d)15%(?!\d)"),                         # 인력 감축률
    re.compile(r"(?<!\d)55%(?!\d)"),                         # 순환자원 대체율 내부 목표치(미공개)
    re.compile(r"H건설"),                                    # 고객사명
    re.compile(r"연\s*12만톤"),                               # 공급계약 물량
    re.compile(r"톤당\s*82,000원"),                           # 공급 단가
    re.compile(r"톤당\s*68,400원"),                           # 제조원가
    re.compile(r"(?<!\d)4\.2%(?!\d)"),                       # 제조원가 상승률
]

PATTERNS = PATTERNS + EXTRA_PATTERNS


def mask(match: re.Match) -> str:
    return "*" * len(match.group())


def mask_text(text: str) -> str:
    for pattern in PATTERNS:
        text = pattern.sub(mask, text)
    return text


def mask_paragraph(paragraph):
    # run 단위로 치환해야 서식(굵기, 색상 등)이 깨지지 않는다.
    for run in paragraph.runs:
        if run.text:
            run.text = mask_text(run.text)


def mask_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                mask_paragraph(paragraph)
            # 표 안에 중첩 표가 있는 경우까지 처리
            for nested_table in cell.tables:
                mask_table(nested_table)


def main():
    document = docx.Document(INPUT_FILE)

    for paragraph in document.paragraphs:
        mask_paragraph(paragraph)

    for table in document.tables:
        mask_table(table)

    # 머리글/바닥글도 함께 처리
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            mask_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            mask_paragraph(paragraph)

    document.save(OUTPUT_FILE)
    print(f"완료: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
